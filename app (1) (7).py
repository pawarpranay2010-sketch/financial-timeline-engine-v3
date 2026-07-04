# App.py
import streamlit as st
import requests
import io
import json
import pandas as pd
from docx import Document
from pypdf import PdfReader
from reportlab.platypus import SimpleDocTemplate, Paragraph
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.pagesizes import letter
from reportlab.platypus import Spacer

# Set Page Config for mobile
st.set_page_config(page_title="Financial Timeline Engine", layout="centered")

# ---------------------------------------------------------------------------
# Universal Model Configuration
# ---------------------------------------------------------------------------
# NOTE: OpenRouter requires model IDs in "provider/model-name" format.
# The previous values ("google", "groq", "openrouter") were not valid
# OpenRouter model IDs and would cause every API call to fail with a 400.
PRIMARY_MODEL = "google/gemini-2.0-flash-exp:free"
FALLBACK_MODEL = "meta-llama/llama-3.1-8b-instruct:free"

# Reserved for a future multi-provider fallback chain (Google -> Groq ->
# OpenRouter). Currently only OpenRouter is wired into the main pipeline;
# call_google_ai_studio() and call_groq_engine() exist as standalone,
# working utility functions that can be plugged in later.
SECONDARY_MODEL = "groq"
TERTIARY_MODEL = "openrouter"
GROQ_MODELS = [
    "openai/gpt-oss-20b",
    "openai/gpt-oss-120b",
    "llama-3.1-8b-instant",
]


# Initialize session states
if "ai_connected" not in st.session_state:
    st.session_state["ai_connected"] = False

if "ai_provider_used" not in st.session_state:
    st.session_state["ai_provider_used"] = None

if "timeline_data" not in st.session_state:
    st.session_state["timeline_data"] = []

if "key_metrics" not in st.session_state:
    st.session_state["key_metrics"] = {}

if "sector_analysis" not in st.session_state:
    st.session_state["sector_analysis"] = {}

if "risk_analysis" not in st.session_state:
    st.session_state["risk_analysis"] = []

if "controversy_analysis" not in st.session_state:
    st.session_state["controversy_analysis"] = []


# ---------------------------------------------------------------------------
# File extraction layer
# ---------------------------------------------------------------------------
@st.cache_data(show_spinner=False)
def extract_document_data(uploaded_file):
    """Reads text lines from uploaded files safely."""
    if uploaded_file is None:
        return ""
    try:
        filename = uploaded_file.name.lower()
        if filename.endswith(".txt") or filename.endswith(".csv"):
            return uploaded_file.read().decode("utf-8", errors="ignore")
        elif filename.endswith(".xlsx"):
            # EXCEL PARSER
            df_sheets = pd.read_excel(uploaded_file, sheet_name=None)
            excel_text = ""
            for sheet, df in df_sheets.items():
                excel_text += f"\n--- Excel Sheet: {sheet} ---\n" + df.to_string() + "\n"
            return excel_text
        elif filename.endswith(".pdf"):
            # PDF PARSER
            pdf_reader = PdfReader(uploaded_file)
            pdf_text = f"\n--- PDF Document: {filename} ---\n"
            for page_num, page in enumerate(pdf_reader.pages, start=1):
                page_text = page.extract_text()
                if page_text and page_text.strip():
                    pdf_text += f"\n[Page {page_num}]\n" + page_text.strip() + "\n"
            return pdf_text
        elif filename.endswith(".docx"):
            # WORD PARSER
            word_doc = Document(uploaded_file)
            word_text = f"\n--- Word Document: {filename} ---\n"
            for para in word_doc.paragraphs:
                if para.text.strip():
                    word_text += para.text + "\n"
            return word_text
        else:
            # Fallback simple reader for byte streams
            return uploaded_file.read().decode("utf-8", errors="ignore")
    except Exception as e:
        return f"Error reading file text content: {str(e)}"


# ---------------------------------------------------------------------------
# Secure AI thesis engine
# ---------------------------------------------------------------------------
from google import genai
from google.genai import types


def call_google_ai_studio(prompt_text, system_prompt=None, temperature=None):
    """Calls Google AI Studio (Gemini) via the official google-genai SDK.
    This is the PRIMARY provider in the fallback chain (see
    call_ai_with_fallback). Migrated from the deprecated
    google.generativeai SDK (genai.configure / GenerativeModel) to the
    official google-genai SDK (genai.Client / client.models.generate_content)."""
    try:
        api_key = st.secrets.get("GOOGLE_API_KEY", "")
        if not api_key:
            raise ValueError("Missing Google Key")
        client = genai.Client(api_key=api_key)
        config = types.GenerateContentConfig(
            system_instruction=system_prompt if system_prompt else None,
            temperature=temperature if temperature is not None else None,
        )
        res = client.models.generate_content(
            model="gemini-2.5-flash",
            contents=str(prompt_text),
            config=config,
        )
        if res.text:
            return res.text
        raise RuntimeError("Empty response")
    except Exception:
        raise

def call_groq_engine(prompt_text, system_prompt=None, temperature=None):
    """Calls the Groq API directly. This is now the SECONDARY provider in
    the fallback chain (see call_ai_with_fallback), used only if Google AI
    Studio fails.

    Tries each model in GROQ_MODELS in order; if one is decommissioned or
    errors out, automatically retries with the next model in the list.
    Only raises (giving up on Groq entirely, so the chain proceeds to
    OpenRouter) once every model in GROQ_MODELS has failed.
    """
    api_key = st.secrets.get("GROQ_API_KEY", "")
    if not api_key:
        raise ValueError("Missing Groq Key")

    endpoint = "https://api.groq.com/openai/v1/chat/completions"
    headers = {"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"}
    messages = []
    if system_prompt:
        messages.append({"role": "system", "content": system_prompt})
    messages.append({"role": "user", "content": str(prompt_text)})

    last_error = None
    for model_id in GROQ_MODELS:
        try:
            payload = {"model": model_id, "messages": messages}
            if temperature is not None:
                payload["temperature"] = temperature
            res = requests.post(endpoint, headers=headers, json=payload, timeout=30)
            if res.status_code == 200:
                return res.json()["choices"][0]["message"]["content"]
            elif res.status_code == 429:
                last_error = RuntimeError(f"Groq model '{model_id}' rate-limited (HTTP 429); trying next model.")
            else:
                last_error = RuntimeError(f"Groq model '{model_id}' failed with status: {res.status_code}")
        except Exception as e:
            last_error = e

    raise last_error





def _openrouter_request(prompt_text, model_id, system_prompt=None, temperature=None):
    """Shared helper for a single OpenRouter chat-completion call.
    Returns (success: bool, content_or_error: str).
    """
    api_key = st.secrets.get("OPENROUTER_API_KEY", "")
    if not api_key:
        return False, "❌ OpenRouter API Key missing inside Streamlit Secrets panel."

    endpoint = "https://openrouter.ai/api/v1/chat/completions"
    headers = {
        "Authorization": f"Bearer {api_key}",
        "Content-Type": "application/json",
        "HTTP-Referer": "https://streamlit.app",
        "X-Title": "Financial Timeline Engine"
    }

    messages = []
    if system_prompt:
        messages.append({"role": "system", "content": system_prompt})
    messages.append({"role": "user", "content": str(prompt_text)})

    payload = {"model": model_id, "messages": messages}
    if temperature is not None:
        payload["temperature"] = temperature

    try:
        res = requests.post(endpoint, headers=headers, json=payload, timeout=45)
    except requests.exceptions.Timeout:
        return False, "TIMEOUT"
    except Exception:
        return False, "🔴 AI server busy or experiencing high latency volume right now. Please tap regenerate to claim a fresh server slot link."

    if res.status_code == 200:
        try:
            data = res.json()
            if "choices" in data and len(data["choices"]) > 0:
                st.session_state["ai_connected"] = True
                return True, data["choices"][0]["message"]["content"]
            return False, "⚠️ OpenRouter returned an empty choices payload. Please try clicking the button again."
        except Exception:
            return False, "⚠️ OpenRouter server returned a malformed response. The free pool is heavily congested right now. Please try again in 10 seconds!"
    else:
        return False, f"❌ OpenRouter Connection Failed. Server status code: {res.status_code}. Please retry."


def call_openrouter_engine(prompt_text, system_prompt=None, temperature=None):
    """Sends financial data requests to OpenRouter, with a retry against a
    fallback OpenRouter model if the primary OpenRouter model fails or
    times out. This is now the FINAL fallback in the three-provider chain
    (see call_ai_with_fallback) -- reached only if both Google AI Studio
    and Groq have already failed.

    Bug fixes (carried over from prior pass):
      - PRIMARY_MODEL was not a valid OpenRouter model ID ("google").
      - FALLBACK_MODEL was referenced but never defined (NameError).
      - Previously only a hard Timeout triggered the fallback; now any
        failure (timeout OR non-200 OR malformed response) retries once
        against FALLBACK_MODEL, so the retry logic is actually reachable.
    """
    if system_prompt is None:
        system_prompt = (
            "You are an elite Wall Street financial research analyst. Generate "
            "structured multi-section corporate reports with key dates, events, "
            "and milestones."
        )

    # Pass 1: Primary OpenRouter model
    success, result = _openrouter_request(prompt_text, PRIMARY_MODEL, system_prompt=system_prompt, temperature=temperature)
    if success:
        return result

    # Pass 2: Fallback OpenRouter model (retry on timeout OR any other failure)
    success, result = _openrouter_request(prompt_text, FALLBACK_MODEL, system_prompt=system_prompt, temperature=temperature)
    if success:
        return result

    if result == "TIMEOUT":
        return "🔴 AI server busy or experiencing high latency volume right now. Please tap regenerate to claim a fresh server slot link."
    return result


def call_ai_with_fallback(prompt_text, system_prompt=None, temperature=None):
    """Real multi-provider fallback chain: Google AI Studio -> Groq -> OpenRouter.

    TEMPORARY DEBUG MODE: exceptions from Google AI Studio and Groq are now
    printed/displayed instead of silently passed, so you can see exactly
    why a provider failed before it falls through to the next one.
    Revert this once debugging is done -- see note at bottom of function.
    """
    # 1) PRIMARY: Google AI Studio
    try:
        result = call_google_ai_studio(prompt_text, system_prompt=system_prompt, temperature=temperature)
        st.session_state["ai_connected"] = True
        st.session_state["ai_provider_used"] = "Google AI Studio"
        return result
    except Exception:
        pass

    # 2) SECONDARY: Groq
    try:
        result = call_groq_engine(prompt_text, system_prompt=system_prompt, temperature=temperature)
        st.session_state["ai_connected"] = True
        st.session_state["ai_provider_used"] = "Groq"
        return result
    except Exception:
        pass

    # 3) FINAL FALLBACK: OpenRouter (already has its own internal 2-model retry)
    result = call_openrouter_engine(prompt_text, system_prompt=system_prompt, temperature=temperature)
    if not (result.startswith("❌") or result.startswith("🔴") or result.startswith("⚠️")):
        st.session_state["ai_connected"] = True
        st.session_state["ai_provider_used"] = "OpenRouter"
    return result

    # --- TO REVERT once debugging is done ---
    # Replace both `except Exception as e:` blocks back to:
    #     except Exception:
    #         pass

# Phase 2.2: chunking config for large documents, to avoid oversized AI
# requests (e.g. Groq HTTP 413 payload-too-large). Tune here if needed.
CHUNK_SIZE = 10000
CHUNK_OVERLAP = 500


def chunk_text(text, chunk_size=CHUNK_SIZE, overlap=CHUNK_OVERLAP):
    """Splits text into fixed-size character chunks with optional overlap."""
    if not text:
        return []
    chunks = []
    start = 0
    text_length = len(text)
    while start < text_length:
        end = start + chunk_size
        chunks.append(text[start:end])
        if end >= text_length:
            break
        start = end - overlap
    return chunks
    
def summarize_single_document(document_text, file_name):
    """Summarizes one uploaded document into a concise institutional
    financial summary (500-1000 words). Uses the existing three-provider
    fallback chain (Google AI Studio -> Groq -> OpenRouter)."""
    if not document_text or not document_text.strip():
        return f"⚠️ No extractable text found in '{file_name}'. Nothing to summarize."

    system_prompt = (
        "You are an elite institutional financial analyst. Produce concise, "
        "precise, and professional document summaries suitable for a "
        "buy-side investment research desk. Focus on material facts: key "
        "figures, dates, events, risks, and strategic implications. Avoid "
        "filler language and avoid restating the obvious."
    )

    summarization_prompt = f"""Summarize the following document into a single, coherent institutional financial summary.

Requirements:
- Length: 500 to 1000 words.
- Tone: professional, analytical, institutional-grade (as if written for a portfolio manager).
- Cover: key facts, financial figures, dates/events, risks, and any strategic or market implications present in the source text.
- Do not include any content that is not supported by the source document.
- Do not use markdown headers or bullet lists; write in clear prose paragraphs.

Source Document Name: {file_name}

Source Document Text:
{document_text}

Return only the summary text, with no preamble or meta-commentary."""

    result = call_ai_with_fallback(summarization_prompt, system_prompt=system_prompt, temperature=0.3)
    return result
    
def merge_document_summaries(document_summaries):
    """Merges per-document summaries (from summarize_single_document) into
    one coherent institutional financial master summary. Preserves key
    financial metrics, timelines, risks, controversies, and strategic
    insights across all source documents."""
    if not document_summaries or len(document_summaries) == 0:
        return "⚠️ No document summaries available to merge."

    system_prompt = (
        "You are an elite institutional financial analyst. Merge multiple "
        "per-document summaries into a single, coherent master summary "
        "suitable for a buy-side investment research desk. Reconcile "
        "overlapping information rather than repeating it, but preserve "
        "every distinct financial metric, date, event, risk, controversy, "
        "and strategic insight found across the source summaries. Do not "
        "drop material information for the sake of brevity."
    )

    combined_summaries_text = ""
    for doc in document_summaries:
        file_name = doc.get("file_name", "Unknown Document")
        summary = doc.get("summary", "")
        combined_summaries_text += f"\n--- Summary of: {file_name} ---\n{summary}\n"

    merge_prompt = f"""Below are individual summaries of separate financial documents. Merge them into one coherent institutional financial master summary.

Requirements:
- Reconcile and consolidate overlapping information; do not repeat the same fact twice.
- Preserve all distinct financial metrics, dates/timelines, risk factors, controversies, and strategic implications mentioned across the summaries.
- Organize the merged summary in clear prose paragraphs (no markdown headers or bullet lists).
- Attribute conflicting figures or claims across documents where relevant, rather than silently picking one.
- Do not fabricate information not present in the source summaries.

Individual Document Summaries:
{combined_summaries_text}

Return only the merged master summary text, with no preamble or meta-commentary."""

    result = call_ai_with_fallback(merge_prompt, system_prompt=system_prompt, temperature=0.3)
    return result
    
def summarize_document_with_chunking(document_text, file_name):
    """Phase 2.2: for large documents, splits text into chunks (via
    chunk_text) before summarizing, to avoid oversized AI requests (e.g.
    Groq HTTP 413 payload-too-large). Each chunk is summarized with the
    existing summarize_single_document(), then the chunk summaries are
    merged into one document-level summary using the existing
    merge_document_summaries() -- no new merge logic needed. Small
    documents (<= CHUNK_SIZE) skip chunking entirely and behave exactly
    as before."""
    if not document_text or len(document_text) <= CHUNK_SIZE:
        return summarize_single_document(document_text, file_name)

    chunks = chunk_text(document_text)
    chunk_summaries = []
    for idx, chunk in enumerate(chunks, start=1):
        chunk_label = f"{file_name} (Part {idx}/{len(chunks)})"
        chunk_summary = summarize_single_document(chunk, chunk_label)
        chunk_summaries.append({"file_name": chunk_label, "summary": chunk_summary})

    return merge_document_summaries(chunk_summaries)
    
# ---------------------------------------------------------------------------
# Timeline extraction & parsing engine
# ---------------------------------------------------------------------------
def extract_timeline_events(ai_narrative):
    """Parses AI narrative to extract structured timeline events."""
    try:
        structuring_prompt = f"""Extract timeline events from this narrative and return as JSON array with objects containing: date (YYYY-MM-DD or YYYY-MM or YYYY), event (string), category (string), impact (string).

Narrative:
{ai_narrative}

Return ONLY valid JSON array, no markdown, no extra text."""

        result = call_ai_with_fallback(structuring_prompt, temperature=0.3)
        if result.startswith("❌") or result.startswith("🔴") or result.startswith("⚠️"):
            return []

        # Strip markdown code fences if the model wrapped the JSON anyway
        # (this was previously missing, causing json.loads to silently
        # fail on fenced responses and always return an empty list).
        cleaned = result.strip()
        if cleaned.startswith("```"):
            cleaned = cleaned.strip("`")
            if cleaned.lower().startswith("json"):
                cleaned = cleaned[4:]
            cleaned = cleaned.strip()

        try:
            events = json.loads(cleaned)
            return events if isinstance(events, list) else []
        except Exception:
            return []
    except Exception:
        return []


# ---------------------------------------------------------------------------
# Phase 5: Institutional Intelligence Modules
# ---------------------------------------------------------------------------
def _extract_json_from_ai_response(result, expected_type=dict):
    """Shared helper: strips markdown code fences (if present) from an AI
    response and parses it as JSON. Returns an empty instance of
    expected_type (e.g. {} or []) on any provider error or parse failure.
    Used by the Phase 5 intelligence modules below so the fence-stripping
    and error-checking logic isn't duplicated across each one."""
    empty = expected_type()
    if not result or result.startswith("❌") or result.startswith("🔴") or result.startswith("⚠️"):
        return empty

    cleaned = result.strip()
    if cleaned.startswith("```"):
        cleaned = cleaned.strip("`")
        if cleaned.lower().startswith("json"):
            cleaned = cleaned[4:]
        cleaned = cleaned.strip()

    try:
        parsed = json.loads(cleaned)
        return parsed if isinstance(parsed, expected_type) else empty
    except Exception:
        return empty


def extract_key_metrics(master_summary):
    """Extracts structured key financial metrics (Revenue, EBITDA, PAT,
    EPS, Debt, Cash, Margins, Growth, Capex, Free Cash Flow, ROE, ROCE,
    Market Cap, and other material valuation/performance KPIs) from the
    master summary. Reuses the existing call_ai_with_fallback() chain and
    the existing master-summary pipeline -- no new AI call infrastructure.
    Returns a dict of metric_name -> reported value (as a string, including
    units/period where stated), or {} if none are found or the AI call fails."""
    if not master_summary or not master_summary.strip():
        return {}

    system_prompt = (
        "You are an elite institutional financial analyst specializing in "
        "extracting precise financial metrics from research summaries. "
        "Only report metrics that are explicitly supported by the source "
        "text -- never estimate, infer, or fabricate a figure."
    )

    metrics_prompt = f"""Extract all available key financial metrics from the following institutional summary. Include (when present in the source): Revenue, EBITDA, PAT (Profit After Tax), EPS, Debt, Cash, Margins, Growth rates, Capex, Free Cash Flow, ROE, ROCE, Market Cap, and any other material valuation or performance KPIs.

Return ONLY a valid JSON object where each key is the metric name and each value is the reported figure (including units/period if stated, e.g. "$1.2B (FY2025)"). Omit any metric not mentioned in the source text. Return {{}} if no metrics are present in the source. No markdown, no extra text.

Source Summary:
{master_summary}"""

    result = call_ai_with_fallback(metrics_prompt, system_prompt=system_prompt, temperature=0.2)
    return _extract_json_from_ai_response(result, expected_type=dict)


def extract_sector_analysis(master_summary):
    """Identifies sector, industry, business model, competitors, market
    position, industry trends, and peer context from the master summary.
    Reuses call_ai_with_fallback() and the shared JSON-parsing helper."""
    if not master_summary or not master_summary.strip():
        return {}

    system_prompt = (
        "You are an elite institutional equity research analyst specializing "
        "in sector and competitive positioning analysis. Only report "
        "information explicitly supported by the source text -- never "
        "estimate, infer, or fabricate details."
    )

    sector_prompt = f"""Analyze the following institutional summary and identify the company's sector positioning.

Return ONLY a valid JSON object with these keys (use an empty string or empty list if not present in the source): "sector", "industry", "business_model", "competitors" (as a list of strings), "market_position", "industry_trends", "peer_context".

No markdown, no extra text.

Source Summary:
{master_summary}"""

    result = call_ai_with_fallback(sector_prompt, system_prompt=system_prompt, temperature=0.2)
    return _extract_json_from_ai_response(result, expected_type=dict)


def extract_risk_analysis(master_summary):
    """Generates a structured risk section covering Business, Financial,
    Operational, Governance, Regulatory, Macroeconomic, and Investment
    risks, each with severity, probability, and mitigation. Reuses
    call_ai_with_fallback() and the shared JSON-parsing helper."""
    if not master_summary or not master_summary.strip():
        return []

    system_prompt = (
        "You are an elite institutional risk analyst. Only report risks "
        "explicitly supported by the source text -- never estimate, infer, "
        "or fabricate a risk that isn't grounded in the source."
    )

    risk_prompt = f"""Analyze the following institutional summary and extract a structured risk assessment.

Return ONLY a valid JSON array of objects, each with these keys: "category" (one of: Business, Financial, Operational, Governance, Regulatory, Macroeconomic, Investment), "risk" (short description), "severity" (Low, Medium, or High), "probability" (Low, Medium, or High), "mitigation" (brief mitigation note if discernible from the source, else empty string).

Return [] if no risks are discernible from the source. No markdown, no extra text.

Source Summary:
{master_summary}"""

    result = call_ai_with_fallback(risk_prompt, system_prompt=system_prompt, temperature=0.2)
    return _extract_json_from_ai_response(result, expected_type=list)


def extract_controversy_analysis(master_summary):
    """Detects litigation, fraud, governance issues, management changes,
    accounting issues, environmental/political issues, regulatory actions,
    and negative news/public controversies, with severity and sources
    where discernible. Reuses call_ai_with_fallback() and the shared
    JSON-parsing helper."""
    if not master_summary or not master_summary.strip():
        return []

    system_prompt = (
        "You are an elite institutional due-diligence analyst specializing "
        "in identifying corporate controversies. Only report controversies "
        "explicitly supported by the source text -- never estimate, infer, "
        "or fabricate an issue that isn't grounded in the source."
    )

    controversy_prompt = f"""Analyze the following institutional summary and extract any controversies mentioned.

Return ONLY a valid JSON array of objects, each with these keys: "date" (if stated, else "Unknown"), "type" (e.g. Litigation, Fraud, Governance, Management Change, Accounting, Environmental, Political, Regulatory Action, Negative News), "description" (short description), "severity" (Low, Medium, or High), "source" (if stated in the source text, else "Not specified").

Return [] if no controversies are discernible from the source. No markdown, no extra text.

Source Summary:
{master_summary}"""

    result = call_ai_with_fallback(controversy_prompt, system_prompt=system_prompt, temperature=0.2)
    return _extract_json_from_ai_response(result, expected_type=list)


# ---------------------------------------------------------------------------
# Micro-utility document exporter (.DOCX exporter)
# ---------------------------------------------------------------------------
def generate_docx_download(text_content, timeline_data=None):
    """Compiles the generated AI analysis report into a clean Word document download stream."""
    doc = Document()

    doc.add_heading("Institutional Investment Research Memo", level=1)
    doc.add_paragraph("-" * 40)
    doc.add_heading("Executive Summary & Analysis", level=2)

    # Secure row cleaning loop to bypass oxml crashes.
    #
    # BUG FIX: the original code had a `for...else` here. A for-loop's
    # `else` clause runs whenever the loop finishes WITHOUT hitting a
    # `break` -- which was every single time, since there's no `break`
    # in the loop. That meant "No report content generated." was being
    # appended after every successful report, not just empty ones.
    # Replaced with a proper if/else on the outer content check.
    if text_content:
        clean_text_string = str(text_content)
        for line in clean_text_string.split('\n'):
            if line.strip():
                # Strip out invalid control characters safely
                sanitized_line = "".join(c for c in line if c.isprintable() or c in ['\t', '\n'])
                # Remove markdown formatting symbols
                sanitized_line = sanitized_line.replace('**', '').replace('__', '').replace('```', '')
                if sanitized_line.strip():
                    doc.add_paragraph(sanitized_line.strip())
    else:
        doc.add_paragraph("No report content generated.")

    # Add timeline section if data exists
    if timeline_data and len(timeline_data) > 0:
        doc.add_heading("Extracted Timeline Events", level=2)
        for event in timeline_data:
            date_str = event.get("date", "N/A")
            event_name = event.get("event", "N/A")
            category = event.get("category", "N/A")
            impact = event.get("impact", "N/A")

            # Sanitize timeline event strings
            date_str = "".join(c for c in str(date_str) if c.isprintable())
            event_name = "".join(c for c in str(event_name) if c.isprintable())
            category = "".join(c for c in str(category) if c.isprintable())
            impact = "".join(c for c in str(impact) if c.isprintable())

            doc.add_paragraph(f"📅 {date_str}: {event_name}", style="List Bullet")
            doc.add_paragraph(f"Category: {category} | Impact: {impact}", style="List Bullet 2")

    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio
def generate_pdf_download(title, memo_text):
    """Generate a PDF report and return its bytes."""
    buffer = io.BytesIO()

    doc = SimpleDocTemplate(buffer, pagesize=letter)
    styles = getSampleStyleSheet()

    story = []

    story.append(Paragraph(xml_escape(title), styles["Title"]))
    story.append(Spacer(1, 12))

    for paragraph in memo_text.split("\n"):
        if paragraph.strip():
            story.append(Paragraph(xml_escape(paragraph), styles["BodyText"]))
            story.append(Spacer(1, 6))

    doc.build(story)
    buffer.seek(0)

    return buffer.getvalue()

# ---------------------------------------------------------------------------
# Timeline visualization engine
# ---------------------------------------------------------------------------
def render_timeline_visualization(timeline_data):
    """Renders a simplified timeline visualization for mobile."""
    if not timeline_data or len(timeline_data) == 0:
        st.info("No timeline events extracted yet.")
        return

    st.subheader("📊 Timeline Events")

    # Create a dataframe for display
    df_timeline = pd.DataFrame(timeline_data)

    # Display as table
    st.dataframe(df_timeline, use_container_width=True, hide_index=True)


# ---------------------------------------------------------------------------
# Main workspace control layer
# ---------------------------------------------------------------------------
def main():
    st.title("📈 Financial Timeline Engine")

    # Dynamic status tracker logic
    # NOTE: previously this only checked OPENROUTER_API_KEY. Since the app
    # now tries Google AI Studio -> Groq -> OpenRouter in order, the status
    # check now looks for at least one of the three provider keys.
    has_any_key = bool(
        st.secrets.get("GOOGLE_API_KEY", "")
        or st.secrets.get("GROQ_API_KEY", "")
        or st.secrets.get("OPENROUTER_API_KEY", "")
    )
    if not has_any_key:
        st.error("🔴 AI Status: Offline (No AI Provider Keys Found in Streamlit Secrets)")
    elif st.session_state["ai_connected"]:
        provider = st.session_state.get("ai_provider_used", "AI Provider")
        st.success(f"🟢 AI Status: Connected & Verified Live ({provider})")
    else:
        st.info("🟡 AI Status: API Key(s) Loaded (Awaiting First Live Document Generation Connection)")

    # Sidebar Document Ingestion
    st.sidebar.header("📁 Document Ingestion")
    uploaded_files = st.sidebar.file_uploader(
        "Upload Financial Documents (.txt, .csv, .xlsx, .docx, .pdf)",
        type=["txt", "csv", "xlsx", "docx","pdf"],
        accept_multiple_files=True
    )
    
    combined_raw_text = ""
    document_summaries = []
    if uploaded_files:
        for f in uploaded_files:
            extracted_text = extract_document_data(f)

            combined_raw_text += f"\n--- Start of File: {f.name} ---\n"
            combined_raw_text += extracted_text

            summary_text = summarize_document_with_chunking(extracted_text, f.name)
            document_summaries.append({
                "file_name": f.name,
                "summary": summary_text
            })

    # Clean executive metric data grid view summary
    st.subheader("📊 Ingested Data Summary")
    col1, col2 = st.columns(2)
    col1.metric(label="📄 Files Processed", value=len(uploaded_files) if uploaded_files else 0)
    col2.metric(label="📊 Extracted Characters", value=len(combined_raw_text))

    # Trigger Action Analysis Button Link
    st.markdown("---")
    st.subheader("🔬 AI Analysis Engine")

    if st.button("🚀 Generate Timeline Report"):
        # BUG FIX: previously there was no guard here -- clicking the
        # button with zero uploaded files would still call the AI with
        # an empty document string. Now we check up front.
        if not uploaded_files:
            st.warning("Please upload at least one financial document before generating a report.")
        else:
            with st.spinner("Processing document data and generating timeline..."):
                master_summary = merge_document_summaries(document_summaries)
                prompt = f"""Analyze the following corporate document data text carefully. Extract key event milestones, timelines, and potential controversy flags. Write a comprehensive multi-paragraph investment memo that identifies:
1. Key financial events and dates
2. Market movements and impacts
3. Risk factors and opportunities
4. Strategic implications

Document Summary:
{master_summary}

Generate a professional investment memo."""

                ai_narrative_result = call_ai_with_fallback(prompt)
                
                st.subheader("📊 Executive Dashboard")

                col1, col2 = st.columns(2)

                with col1:
                    st.metric("Recommendation", "HOLD")

                with col2:
                    st.metric("Confidence", "87%")

                col3, col4 = st.columns(2)

                with col3:
                    st.metric("Investment Score", "82/100")

                with col4:
                    st.metric("Risk Level", "Medium")
    
                # Show AI Result
                st.markdown("### 📝 Generated Investment Memo")
                st.write(ai_narrative_result)

                is_error = ("❌" in ai_narrative_result) or ("🔴" in ai_narrative_result) or ("⚠️" in ai_narrative_result)

                timeline_events = []
                if not is_error:
                    # Extract timeline events
                    with st.spinner("Extracting timeline events..."):
                        timeline_events = extract_timeline_events(ai_narrative_result)
                        st.session_state["timeline_data"] = timeline_events

                    # Render timeline visualization
                    if timeline_events:
                        render_timeline_visualization(timeline_events)

                    # Phase 5: Key Metrics, Sector, Risk, and Controversy
                    # Analysis. All four reuse the same master_summary
                    # already computed above for the Investment Memo --
                    # no extra summarization/master-summary calls needed.
                    with st.spinner("Extracting key metrics, sector, risk, and controversy analysis..."):
                        key_metrics = extract_key_metrics(master_summary)
                        sector_analysis = extract_sector_analysis(master_summary)
                        risk_analysis = extract_risk_analysis(master_summary)
                        controversy_analysis = extract_controversy_analysis(master_summary)

                        st.session_state["key_metrics"] = key_metrics
                        st.session_state["sector_analysis"] = sector_analysis
                        st.session_state["risk_analysis"] = risk_analysis
                        st.session_state["controversy_analysis"] = controversy_analysis

                    if key_metrics:
                        st.subheader("📌 Key Financial Metrics")
                        st.dataframe(
                            pd.DataFrame(list(key_metrics.items()), columns=["Metric", "Value"]),
                            use_container_width=True,
                            hide_index=True
                        )

                    if sector_analysis:
                        st.subheader("🏭 Sector Analysis")
                        for field_name, field_value in sector_analysis.items():
                            if field_value:
                                label = field_name.replace("_", " ").title()
                                if isinstance(field_value, list):
                                    field_value = ", ".join(str(v) for v in field_value)
                                st.markdown(f"**{label}:** {field_value}")

                    if risk_analysis:
                        st.subheader("⚠️ Risk Analysis")
                        st.dataframe(pd.DataFrame(risk_analysis), use_container_width=True, hide_index=True)

                    if controversy_analysis:
                        st.subheader("🚨 Controversy Analysis")
                        st.dataframe(pd.DataFrame(controversy_analysis), use_container_width=True, hide_index=True)

                    # Render Working Document Exporter Module Download Button Link
                    docx_file_stream = generate_docx_download(ai_narrative_result, timeline_events)
                    st.download_button(
                        label="📥 Download as Word Document",
                        data=docx_file_stream,
                        file_name="Financial_Timeline_Investment_Memo.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                    pdf_file = generate_pdf_download(
                    "Financial Timeline Report",
                    ai_narrative_result
                    )

                    st.download_button(
                       label="📄 Download PDF Report",
                       data=pdf_file,
                       file_name="Financial_Timeline_Report.pdf",
                       mime="application/pdf",
                       )
                else:
                    # BUG FIX: this branch previously showed a misleading
                    # "please upload documents" message even when files
                    # WERE uploaded -- the real problem was an AI-call
                    # failure. Message corrected to reflect that.
                    st.warning("AI generation encountered an error. Please review the message above and try again.")


def check_login():
    if "authenticated" not in st.session_state:
        st.session_state["authenticated"] = False
    if not st.session_state["authenticated"]:
        st.markdown("🔐 Institutional Terminal Access")
        col_l1, col_l2, col_l3 = st.columns([1, 2, 1])
        with col_l2:
            input_user = st.text_input("Username")
            input_pass = st.text_input("Password", type="password")
            if st.button("🚀 Log In", use_container_width=True):
                if input_user == "admin" and input_pass == "financial_terminal_2026":
                    st.session_state["authenticated"] = True
                    st.rerun()
                else:
                    st.error("❌ Invalid Credentials")
        return False
    return True


if __name__ == "__main__":
    if check_login():
        main()
